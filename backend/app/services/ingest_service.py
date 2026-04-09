from __future__ import annotations

import csv
import io
import json
import re
from collections import Counter
from dataclasses import dataclass
from html import unescape
from pathlib import Path
from urllib.parse import parse_qs, urlparse
from uuid import uuid4

import pandas as pd
import requests
from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

from app.core.config import settings
from app.services.retrieval_service import retrieval_service
from app.services.storage_service import storage_service


SAFE_NAME_RE = re.compile(r"[^A-Za-z0-9._-]+")
CATALOG_SOURCE_ID = "catalog"
CATALOG_SOURCE_LABEL = "Catalog (All PDF Sources)"
COURSE_CODE_RE = re.compile(r"\b[A-Z]{2,5}\s?\d{3}\b")
WORD_RE = re.compile(r"[A-Za-z][A-Za-z0-9_-]{1,24}")
HEADING_NUMBER_RE = re.compile(r"^(?:\d+(?:\.\d+){0,3}|[IVXLCM]+)\s+")
LIST_LINE_RE = re.compile(r"^\s*(?:[-*\u2022]|(?:\d+[\.)]))\s+")
TABLE_SPACING_RE = re.compile(r"\S+\s{2,}\S+")
ISO_DATE_RE = re.compile(r"\b\d{4}-\d{2}-\d{2}\b")
MONTH_DATE_RE = re.compile(
    r"\b(?:jan|january|feb|february|mar|march|apr|april|may|jun|june|jul|july|aug|august|sep|sept|september|"
    r"oct|october|nov|november|dec|december)\s+\d{1,2}(?:,\s*\d{4})?\b",
    re.I,
)
TIME_RE = re.compile(r"\b\d{1,2}:\d{2}\s?(?:AM|PM|am|pm)?\b")
ROOM_RE = re.compile(
    r"\b(?:room|hall|building|campus)\s*[A-Za-z0-9-]+\b|\b(?:online|zoom)\b|\blab\s*[A-Z0-9-]{1,8}\b",
    re.I,
)
WEEK_RE = re.compile(r"\bweek\s+\d+\b", re.I)
STOPWORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "for",
    "from",
    "in",
    "into",
    "is",
    "it",
    "of",
    "on",
    "or",
    "that",
    "the",
    "this",
    "to",
    "with",
}
HTML_SCRIPT_RE = re.compile(r"<(script|style)[^>]*>.*?</\1>", re.I | re.S)
HTML_TAG_RE = re.compile(r"<[^>]+>")
GOOGLE_SHEETS_HOSTS = {"docs.google.com"}
GOOGLE_SHEET_ID_RE = re.compile(r"/spreadsheets/(?:u/\d+/)?d/([a-zA-Z0-9-_]+)")
GID_RE = re.compile(r"(?:^|[&#])gid=(\d+)")
CANONICAL_COLUMN_LABELS: dict[str, tuple[str, ...]] = {
    "course_code": ("course code", "course no", "course number", "code", "course"),
    "course_name": ("course name", "subject", "title", "module"),
    "exam_date": ("exam date", "date", "test date"),
    "exam_time": ("exam time", "time", "start time"),
    "location": ("location", "room", "venue", "hall"),
    "notes": ("note", "notes", "remark", "remarks", "comment"),
}


@dataclass
class ParsedSection:
    section_title: str
    parent_section_title: str | None
    page_start: int | None
    page_end: int | None
    section_type: str
    section_text: str
    keywords: list[str]
    facts: dict[str, list[str]]


@dataclass
class ParsedDocument:
    transcript: str
    sections: list[ParsedSection]


@dataclass
class SheetTable:
    sheet_name: str
    headers: list[str]
    rows: list[list[str]]


def safe_filename(name: str) -> str:
    cleaned = SAFE_NAME_RE.sub("_", name).strip("._")
    return cleaned or "upload.bin"


def normalize_line(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "")).strip()


def unique_values(values: list[str]) -> list[str]:
    seen: set[str] = set()
    output: list[str] = []
    for value in values:
        clean = normalize_line(value)
        if not clean:
            continue
        key = clean.lower()
        if key in seen:
            continue
        output.append(clean)
        seen.add(key)
    return output


def extract_section_facts(title: str, text: str) -> dict[str, list[str]]:
    probe = f"{title}\n{text}"
    return {
        "dates": unique_values(ISO_DATE_RE.findall(probe) + MONTH_DATE_RE.findall(probe)),
        "times": unique_values(TIME_RE.findall(probe)),
        "locations": unique_values(ROOM_RE.findall(probe)),
        "weeks": unique_values(WEEK_RE.findall(probe)),
        "course_codes": unique_values([re.sub(r"\s+", " ", code).strip() for code in COURSE_CODE_RE.findall(probe)]),
    }


def read_pdf_pages(path: Path) -> list[tuple[int, str]]:
    reader = PdfReader(str(path))
    pages: list[tuple[int, str]] = []
    for idx, page in enumerate(reader.pages, start=1):
        pages.append((idx, page.extract_text() or ""))
    return pages


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def read_csv(path: Path) -> str:
    rows: list[str] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rows)


def read_xlsx(path: Path) -> str:
    lines: list[str] = []
    excel = pd.ExcelFile(path)
    for sheet in excel.sheet_names[:6]:
        df = excel.parse(sheet_name=sheet)
        lines.append(f"[Sheet: {sheet}]")
        lines.append(df.fillna("").to_csv(index=False))
    return "\n".join(lines)


def is_heading(line: str) -> tuple[bool, int]:
    clean = normalize_line(line)
    if not clean or len(clean) > 120:
        return False, 0

    words = re.findall(r"[A-Za-z][A-Za-z0-9/&-]*", clean)
    if not words or len(words) > 16:
        return False, 0

    if HEADING_NUMBER_RE.match(clean):
        level = clean.split(" ", 1)[0].count(".") + 1
        return True, max(1, min(level, 4))

    alpha = [char for char in clean if char.isalpha()]
    upper_ratio = (sum(1 for char in alpha if char.isupper()) / len(alpha)) if alpha else 0.0
    if upper_ratio >= 0.85 and len(words) <= 12:
        return True, 1

    if clean == clean.title() and len(words) <= 10 and not clean.endswith((".", "?", "!")):
        return True, 2

    return False, 0


def is_table_or_list_line(line: str) -> bool:
    clean = normalize_line(line)
    if not clean:
        return False
    if LIST_LINE_RE.match(clean):
        return True
    if "|" in clean or "\t" in line:
        return True
    if TABLE_SPACING_RE.search(line):
        return True
    if clean.count(",") >= 3:
        return True
    return False


def classify_section(title: str, text: str, table_like_ratio: float) -> str:
    probe = f"{title} {text[:2000]}".lower()
    if any(token in probe for token in ("study plan", "course no", "course code", "curriculum", "corequisite", "prerequisite")):
        return "courses"
    if any(token in probe for token in ("requirement", "admission", "eligibility", "must complete", "credit hour")):
        return "requirements"
    if any(token in probe for token in ("schedule", "calendar", "semester", "term", "week")):
        return "schedule"
    if any(token in probe for token in ("policy", "regulation", "guideline", "conduct", "attendance")):
        return "policy"
    if any(token in probe for token in ("program", "major", "bachelor", "master", "degree")):
        return "program"
    if table_like_ratio >= 0.2:
        return "table_or_list"
    return "general"


def extract_keywords(title: str, text: str) -> list[str]:
    tokens = [token.lower() for token in WORD_RE.findall(f"{title} {text}") if token.lower() not in STOPWORDS]
    top = [value for value, _ in Counter(tokens).most_common(16)]
    courses = [re.sub(r"\s+", " ", code).strip() for code in COURSE_CODE_RE.findall(text)]
    seen: set[str] = set()
    output: list[str] = []
    for value in [*courses, *top]:
        key = value.lower()
        if key in seen:
            continue
        output.append(value)
        seen.add(key)
        if len(output) >= 20:
            break
    return output


def build_document_map(pages: list[tuple[int | None, str]], fallback_title: str) -> ParsedDocument:
    section_stack: list[tuple[int, str]] = []
    sections: list[ParsedSection] = []
    transcript_parts: list[str] = []

    current_title = fallback_title
    current_parent: str | None = None
    current_lines: list[str] = []
    current_page_start: int | None = None
    current_page_end: int | None = None
    table_like_count = 0
    content_line_count = 0

    def flush_section() -> None:
        nonlocal current_title, current_parent, current_lines
        nonlocal current_page_start, current_page_end, table_like_count, content_line_count
        text = "\n".join(line for line in current_lines if line).strip()
        if not text:
            current_lines = []
            current_page_start = None
            current_page_end = None
            table_like_count = 0
            content_line_count = 0
            return
        ratio = (table_like_count / content_line_count) if content_line_count else 0.0
        sections.append(
            ParsedSection(
                section_title=current_title or fallback_title,
                parent_section_title=current_parent,
                page_start=current_page_start,
                page_end=current_page_end or current_page_start,
                section_type=classify_section(current_title, text, ratio),
                section_text=text,
                keywords=extract_keywords(current_title, text),
                facts=extract_section_facts(current_title, text),
            )
        )
        current_lines = []
        current_page_start = None
        current_page_end = None
        table_like_count = 0
        content_line_count = 0

    for page_number, raw_text in pages:
        page_text = raw_text or ""
        if page_number is not None:
            transcript_parts.append(f"[Page {page_number}]")
        page_lines = page_text.splitlines() if page_text else [""]
        for line in page_lines:
            clean = normalize_line(line)
            if not clean:
                continue

            heading, level = is_heading(clean)
            if heading:
                flush_section()
                while section_stack and section_stack[-1][0] >= level:
                    section_stack.pop()
                current_parent = section_stack[-1][1] if section_stack else None
                current_title = clean
                section_stack.append((level, clean))
                current_lines = [clean]
                if page_number is not None:
                    current_page_start = page_number
                    current_page_end = page_number
                continue

            current_lines.append(clean)
            content_line_count += 1
            if is_table_or_list_line(line):
                table_like_count += 1
            if page_number is not None:
                if current_page_start is None:
                    current_page_start = page_number
                current_page_end = page_number

        normalized_page = "\n".join(normalize_line(value) for value in page_lines if normalize_line(value))
        if normalized_page:
            transcript_parts.append(normalized_page)

    flush_section()

    if not sections:
        joined = "\n\n".join(part for _, part in pages if part.strip()).strip()
        if joined:
            sections = [
                ParsedSection(
                    section_title=fallback_title,
                    parent_section_title=None,
                    page_start=pages[0][0] if pages else None,
                    page_end=pages[-1][0] if pages else None,
                    section_type="general",
                    section_text=joined,
                    keywords=extract_keywords(fallback_title, joined),
                    facts=extract_section_facts(fallback_title, joined),
                )
            ]

    transcript = "\n\n".join(transcript_parts).strip()
    if not transcript:
        transcript = "\n\n".join(section.section_text for section in sections).strip()
    return ParsedDocument(transcript=transcript, sections=sections)


def html_to_text(value: str) -> str:
    no_script = HTML_SCRIPT_RE.sub(" ", value or "")
    no_tags = HTML_TAG_RE.sub("\n", no_script)
    cleaned = unescape(no_tags)
    lines = [normalize_line(line) for line in cleaned.splitlines()]
    return "\n".join(line for line in lines if line)


def safe_url_basename(url: str) -> str:
    parsed = urlparse(url)
    candidate = Path(parsed.path).name or "source"
    if "." not in candidate:
        return f"{candidate}.txt"
    return candidate


def _normalize_header_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (value or "").strip().lower()).strip()


def _clean_cell(value: object) -> str:
    if value is None:
        return ""
    text = normalize_line(str(value))
    if text.lower() == "nan":
        return ""
    return text


def _guess_canonical_fields(headers: list[str]) -> dict[int, str]:
    mappings: dict[int, str] = {}
    for idx, raw in enumerate(headers):
        normalized = _normalize_header_key(raw)
        if not normalized:
            continue
        for canonical, labels in CANONICAL_COLUMN_LABELS.items():
            if any(label in normalized for label in labels):
                mappings[idx] = canonical
                break
    return mappings


def _safe_sheet_name(value: str, fallback_index: int) -> str:
    clean = normalize_line(value)
    if not clean:
        return f"Sheet {fallback_index + 1}"
    return clean[:120]


def _build_schedule_document_from_tables(tables: list[SheetTable], fallback_title: str) -> ParsedDocument:
    sections: list[ParsedSection] = []
    transcript_blocks: list[str] = []

    for idx, table in enumerate(tables):
        sheet_name = _safe_sheet_name(table.sheet_name, idx)
        headers = [normalize_line(cell) or f"Column {col_idx + 1}" for col_idx, cell in enumerate(table.headers)]
        canonical_fields = _guess_canonical_fields(headers)
        row_lines: list[str] = []
        max_rows = min(len(table.rows), 800)

        for row_index in range(max_rows):
            row = table.rows[row_index]
            if not row:
                continue
            normalized_row = [_clean_cell(value) for value in row]
            if not any(normalized_row):
                continue

            mapped: dict[str, str] = {}
            extras: list[str] = []
            for col_idx, value in enumerate(normalized_row):
                if not value:
                    continue
                canonical = canonical_fields.get(col_idx)
                header_name = headers[col_idx] if col_idx < len(headers) else f"Column {col_idx + 1}"
                if canonical and canonical not in mapped:
                    mapped[canonical] = value
                else:
                    extras.append(f"{header_name}: {value}")

            canonical_parts: list[str] = []
            ordered = [
                ("course_code", "Course Code"),
                ("course_name", "Course Name"),
                ("exam_date", "Exam Date"),
                ("exam_time", "Exam Time"),
                ("location", "Location"),
                ("notes", "Notes"),
            ]
            for key, label in ordered:
                value = mapped.get(key)
                if value:
                    canonical_parts.append(f"{label}: {value}")

            if not canonical_parts and not extras:
                continue

            joined_parts = "; ".join([*canonical_parts, *extras])
            row_lines.append(f"Exam schedule row {row_index + 1}: {joined_parts}")

        if not row_lines:
            continue

        detected_fields = unique_values([field for field in canonical_fields.values()])
        section_text = "\n".join(
            [
                f"Sheet: {sheet_name}",
                (
                    "Detected fields: "
                    + ", ".join(detected_fields)
                    if detected_fields
                    else "Detected fields: none"
                ),
                *row_lines,
            ]
        ).strip()

        section = ParsedSection(
            section_title=f"Sheet: {sheet_name}",
            parent_section_title=fallback_title,
            page_start=None,
            page_end=None,
            section_type="schedule",
            section_text=section_text,
            keywords=extract_keywords(sheet_name, section_text),
            facts=extract_section_facts(sheet_name, section_text),
        )
        sections.append(section)
        transcript_blocks.append(section_text)

    transcript = "\n\n".join(transcript_blocks).strip()
    if not sections:
        return ParsedDocument(
            transcript=transcript,
            sections=[],
        )
    return ParsedDocument(
        transcript=transcript,
        sections=sections,
    )


def _extract_google_sheet_identifiers(parsed_url) -> tuple[str | None, str | None]:
    sheet_id_match = GOOGLE_SHEET_ID_RE.search(parsed_url.path or "")
    spreadsheet_id = sheet_id_match.group(1) if sheet_id_match else None

    gid: str | None = None
    query_values = parse_qs(parsed_url.query or "")
    if query_values.get("gid"):
        gid = normalize_line(query_values["gid"][0])
    if not gid:
        fragment_gid = GID_RE.search(parsed_url.fragment or "")
        if fragment_gid:
            gid = fragment_gid.group(1)

    return spreadsheet_id, gid


def _is_google_sheets_url(parsed_url) -> bool:
    host = (parsed_url.netloc or "").lower()
    return host in GOOGLE_SHEETS_HOSTS and "/spreadsheets/" in (parsed_url.path or "")


def extract_document(path: Path) -> ParsedDocument:
    ext = path.suffix.lower()
    if ext == ".pdf":
        pages = read_pdf_pages(path)
        return build_document_map([(page, text) for page, text in pages], fallback_title=path.stem or "Document")
    if ext in {".txt", ".md", ".log"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return build_document_map([(None, text)], fallback_title=path.stem or "Document")
    if ext == ".json":
        raw = path.read_text(encoding="utf-8", errors="ignore")
        try:
            parsed = json.loads(raw)
            text = json.dumps(parsed, ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            text = raw
        return build_document_map([(None, text)], fallback_title=path.stem or "Document")
    if ext == ".docx":
        return build_document_map([(None, read_docx(path))], fallback_title=path.stem or "Document")
    if ext == ".csv":
        return build_document_map([(None, read_csv(path))], fallback_title=path.stem or "Document")
    if ext == ".xlsx":
        return build_document_map([(None, read_xlsx(path))], fallback_title=path.stem or "Document")
    raise ValueError(f"Unsupported file type: {ext}")


class IngestService:
    def _request_url(self, url: str, timeout: int) -> requests.Response:
        headers = {"User-Agent": "Local-Academic-AI-Ingest/1.0"}
        try:
            return requests.get(url, timeout=timeout, headers=headers)
        except requests.RequestException as exc:
            raise ValueError(f"Could not fetch URL: {exc}") from exc

    def _sheet_tables_from_xlsx_bytes(self, payload: bytes) -> list[SheetTable]:
        if not payload:
            return []
        excel = pd.ExcelFile(io.BytesIO(payload))
        tables: list[SheetTable] = []
        for sheet_name in excel.sheet_names[:12]:
            df = excel.parse(sheet_name=sheet_name).fillna("")
            headers = [_clean_cell(value) or f"Column {idx + 1}" for idx, value in enumerate(df.columns.tolist())]
            rows = [[_clean_cell(cell) for cell in row] for row in df.values.tolist()]
            tables.append(SheetTable(sheet_name=sheet_name, headers=headers, rows=rows))
        return tables

    def _sheet_tables_from_csv_text(self, text: str, sheet_name: str) -> list[SheetTable]:
        rows: list[list[str]] = []
        reader = csv.reader(io.StringIO(text))
        for row in reader:
            rows.append([_clean_cell(cell) for cell in row])
        if not rows:
            return []
        headers = rows[0]
        body = rows[1:] if len(rows) > 1 else []
        return [SheetTable(sheet_name=sheet_name, headers=headers, rows=body)]

    def _fetch_google_public_sheet_tables(self, spreadsheet_id: str, gid: str | None) -> tuple[list[SheetTable], bytes | None, str]:
        timeout = settings.url_ingest_timeout_seconds
        xlsx_url = f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=xlsx"
        xlsx_response = self._request_url(xlsx_url, timeout=timeout)
        if xlsx_response.status_code < 400 and xlsx_response.content[:2] == b"PK":
            tables = self._sheet_tables_from_xlsx_bytes(xlsx_response.content)
            if tables:
                return tables, xlsx_response.content, "xlsx"

        if gid:
            csv_url = f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=csv&gid={gid}"
            csv_response = self._request_url(csv_url, timeout=timeout)
            if csv_response.status_code < 400 and csv_response.text.strip():
                tables = self._sheet_tables_from_csv_text(csv_response.text, sheet_name=f"Sheet gid={gid}")
                if tables:
                    return tables, csv_response.content, "csv"
        return [], None, ""

    def _fetch_google_private_sheet_tables(self, spreadsheet_id: str, gid: str | None) -> tuple[list[SheetTable], dict]:
        try:
            from google.oauth2 import service_account
            from googleapiclient.discovery import build
            from googleapiclient.errors import HttpError
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ValueError(
                "Google Sheets private mode requires google-api-python-client and google-auth in backend dependencies."
            ) from exc

        scopes = ["https://www.googleapis.com/auth/spreadsheets.readonly"]
        credentials = None
        if settings.google_sheets_service_account_json.strip():
            try:
                info = json.loads(settings.google_sheets_service_account_json)
            except json.JSONDecodeError as exc:
                raise ValueError("GOOGLE_SHEETS_SERVICE_ACCOUNT_JSON is not valid JSON.") from exc
            credentials = service_account.Credentials.from_service_account_info(info, scopes=scopes)
        elif settings.google_sheets_service_account_file.strip():
            credentials_path = settings.resolve_path(settings.google_sheets_service_account_file)
            if not credentials_path.exists():
                raise ValueError("GOOGLE_SHEETS_SERVICE_ACCOUNT_FILE was set but the file does not exist.")
            credentials = service_account.Credentials.from_service_account_file(str(credentials_path), scopes=scopes)
        else:
            raise ValueError(
                "Google Sheets link appears private. Configure GOOGLE_SHEETS_SERVICE_ACCOUNT_FILE or GOOGLE_SHEETS_SERVICE_ACCOUNT_JSON for private ingestion."
            )

        service = build("sheets", "v4", credentials=credentials, cache_discovery=False)
        try:
            metadata = (
                service.spreadsheets()
                .get(spreadsheetId=spreadsheet_id)
                .execute(num_retries=1)
            )
        except HttpError as exc:
            raise ValueError(f"Google Sheets API could not read this sheet: {exc}") from exc

        spreadsheet_title = normalize_line(metadata.get("properties", {}).get("title", "Google Sheet")) or "Google Sheet"
        sheets_meta = metadata.get("sheets", []) or []
        selected_tabs: list[dict] = []

        if gid:
            selected_tabs = [
                entry for entry in sheets_meta if str(entry.get("properties", {}).get("sheetId")) == str(gid)
            ]
            if not selected_tabs:
                selected_tabs = sheets_meta[:1]
        else:
            selected_tabs = sheets_meta[:8]

        tables: list[SheetTable] = []
        for entry in selected_tabs:
            tab_name = normalize_line(entry.get("properties", {}).get("title", "Sheet")) or "Sheet"
            range_name = f"'{tab_name}'"
            values_response = (
                service.spreadsheets()
                .values()
                .get(spreadsheetId=spreadsheet_id, range=range_name)
                .execute(num_retries=1)
            )
            values = values_response.get("values", []) or []
            if not values:
                continue
            headers = [_clean_cell(cell) or f"Column {idx + 1}" for idx, cell in enumerate(values[0])]
            rows = [[_clean_cell(cell) for cell in row] for row in values[1:]]
            tables.append(SheetTable(sheet_name=tab_name, headers=headers, rows=rows))

        return tables, {"title": spreadsheet_title}

    def _ingest_google_sheet_url(self, parsed_url) -> dict:
        spreadsheet_id, gid = _extract_google_sheet_identifiers(parsed_url)
        if not spreadsheet_id:
            raise ValueError("Invalid Google Sheets URL. Expected a /spreadsheets/d/<spreadsheet_id> link.")

        tables, raw_payload, payload_format = self._fetch_google_public_sheet_tables(spreadsheet_id, gid)
        ingest_mode = "google_public"
        ingest_note = "Google Sheet ingested via public export."
        sheet_title = "Google Sheet"
        if tables:
            sheet_title = _safe_sheet_name(tables[0].sheet_name, 0)

        if not tables:
            if not settings.google_sheets_credentials_configured:
                raise ValueError(
                    "This Google Sheet could not be read publicly. It may be private. Configure "
                    "GOOGLE_SHEETS_SERVICE_ACCOUNT_FILE or GOOGLE_SHEETS_SERVICE_ACCOUNT_JSON to enable private ingestion."
                )
            tables, private_meta = self._fetch_google_private_sheet_tables(spreadsheet_id, gid)
            ingest_mode = "google_private_api"
            ingest_note = "Google Sheet ingested via private Google Sheets API."
            sheet_title = normalize_line(private_meta.get("title", "Google Sheet")) or "Google Sheet"

        parsed_doc = _build_schedule_document_from_tables(tables, fallback_title=sheet_title)
        if not parsed_doc.sections or not parsed_doc.transcript.strip():
            raise ValueError("Google Sheet was readable, but no usable schedule rows were found.")

        source_id = f"src_{uuid4().hex}"
        target_path: Path
        if raw_payload and payload_format == "xlsx":
            target_path = settings.uploads_dir_path / safe_filename(f"{source_id}_google_sheet_{spreadsheet_id}.xlsx")
            target_path.write_bytes(raw_payload)
        elif raw_payload and payload_format == "csv":
            suffix = f"_{gid}" if gid else ""
            target_path = settings.uploads_dir_path / safe_filename(f"{source_id}_google_sheet{suffix}.csv")
            target_path.write_bytes(raw_payload)
        else:
            serialized = []
            for table in tables:
                serialized.append(
                    {
                        "sheet_name": table.sheet_name,
                        "headers": table.headers,
                        "rows": table.rows,
                    }
                )
            target_path = settings.uploads_dir_path / safe_filename(f"{source_id}_google_sheet_{spreadsheet_id}.json")
            target_path.write_text(json.dumps(serialized, ensure_ascii=False, indent=2), encoding="utf-8")

        source_name = f"Google Sheet: {sheet_title}" if sheet_title else "Google Sheet"
        storage_service.upsert_source(
            source_id=source_id,
            name=source_name,
            kind="upload",
            file_path=str(target_path),
        )
        chunk_count = retrieval_service.index_source_document(
            source_id=source_id,
            transcript=parsed_doc.transcript,
            sections=parsed_doc.sections,
        )
        return {
            "source_id": source_id,
            "name": source_name,
            "kind": "upload",
            "chunks_indexed": chunk_count,
            "ingest_mode": ingest_mode,
            "ingest_note": ingest_note,
        }

    def ensure_catalog(self) -> None:
        existing = storage_service.get_source(CATALOG_SOURCE_ID)
        if existing and storage_service.get_chunks(CATALOG_SOURCE_ID):
            return

        pdf_files = sorted(settings.catalog_dir_path.glob("*.pdf"))
        if not pdf_files:
            return

        combined_sections: list[ParsedSection] = []
        transcript_parts: list[str] = []
        for pdf in pdf_files:
            parsed = extract_document(pdf)
            if not parsed.transcript.strip():
                continue
            transcript_parts.append(f"[Catalog File: {pdf.name}]\n{parsed.transcript}")
            combined_sections.extend(parsed.sections)

        combined_transcript = "\n\n".join(transcript_parts).strip()
        if not combined_transcript:
            return

        storage_service.upsert_source(
            source_id=CATALOG_SOURCE_ID,
            name=CATALOG_SOURCE_LABEL,
            kind="catalog",
            file_path=str(settings.catalog_dir_path),
        )
        retrieval_service.index_source_document(
            source_id=CATALOG_SOURCE_ID,
            transcript=combined_transcript,
            sections=combined_sections,
        )

    async def ingest_upload(self, file: UploadFile) -> dict:
        original = safe_filename(file.filename or "upload.bin")
        source_id = f"src_{uuid4().hex}"
        target_name = f"{source_id}_{original}"
        target_path = settings.uploads_dir_path / target_name

        raw = await file.read()
        if len(raw) > settings.upload_max_bytes:
            raise ValueError("Uploaded file is too large.")

        target_path.write_bytes(raw)
        parsed = extract_document(target_path)
        if not parsed.transcript.strip():
            raise ValueError("No readable text found in uploaded file.")

        storage_service.upsert_source(
            source_id=source_id,
            name=original,
            kind="upload",
            file_path=str(target_path),
        )
        chunk_count = retrieval_service.index_source_document(
            source_id=source_id,
            transcript=parsed.transcript,
            sections=parsed.sections,
        )

        return {
            "source_id": source_id,
            "name": original,
            "kind": "upload",
            "chunks_indexed": chunk_count,
        }

    def ingest_url(self, url: str) -> dict:
        parsed = urlparse((url or "").strip())
        if parsed.scheme not in {"http", "https"}:
            raise ValueError("Only http/https URLs are supported.")
        if not parsed.netloc:
            raise ValueError("Invalid URL.")
        if _is_google_sheets_url(parsed):
            return self._ingest_google_sheet_url(parsed)

        response = self._request_url(parsed.geturl(), timeout=settings.url_ingest_timeout_seconds)
        if response.status_code >= 400:
            raise ValueError(f"URL fetch failed with status {response.status_code}.")

        content_type = (response.headers.get("content-type") or "").lower()
        source_id = f"src_{uuid4().hex}"
        default_name = safe_filename(safe_url_basename(parsed.geturl()))
        ingest_mode = "url_generic"
        ingest_note = "URL content was ingested successfully."

        if "text/html" in content_type:
            text = html_to_text(response.text)
            if not text.strip():
                raise ValueError("No readable text could be extracted from the page.")
            text = text[: settings.url_ingest_max_chars]
            target_name = safe_filename(f"{source_id}_{Path(default_name).stem}.txt")
            target_path = settings.uploads_dir_path / target_name
            target_path.write_text(text, encoding="utf-8")
            parsed_doc = build_document_map([(None, text)], fallback_title=Path(default_name).stem or "Web Source")
            source_name = f"Web: {parsed.netloc}"
            ingest_mode = "url_public_html"
        else:
            extension = ".pdf" if "application/pdf" in content_type else Path(default_name).suffix.lower() or ".txt"
            target_name = safe_filename(f"{source_id}_{Path(default_name).stem}{extension}")
            target_path = settings.uploads_dir_path / target_name
            if len(response.content) > settings.upload_max_bytes:
                raise ValueError("URL content is too large.")
            target_path.write_bytes(response.content)
            try:
                parsed_doc = extract_document(target_path)
            except ValueError:
                text_payload = response.text[: settings.url_ingest_max_chars]
                target_path.write_text(text_payload, encoding="utf-8")
                parsed_doc = build_document_map([(None, text_payload)], fallback_title=Path(default_name).stem or "Web Source")
            source_name = f"URL: {default_name}"
            ingest_mode = "url_public_file"

        if not parsed_doc.transcript.strip():
            raise ValueError("No readable text found from this URL.")

        storage_service.upsert_source(
            source_id=source_id,
            name=source_name,
            kind="upload",
            file_path=str(target_path),
        )
        chunk_count = retrieval_service.index_source_document(
            source_id=source_id,
            transcript=parsed_doc.transcript,
            sections=parsed_doc.sections,
        )

        return {
            "source_id": source_id,
            "name": source_name,
            "kind": "upload",
            "chunks_indexed": chunk_count,
            "ingest_mode": ingest_mode,
            "ingest_note": ingest_note,
        }


ingest_service = IngestService()
